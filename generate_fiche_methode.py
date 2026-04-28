"""Génère la fiche méthode Word pour le déploiement de l'API Sellsy dans SuiviPDP."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

DOC_PATH = r"C:\Users\ThibaultHOCEDEZ\Documents\suivipdp\FM-IT-XX_Deploiement_API_Sellsy_SuiviPDP.docx"

GREEN = RGBColor(0x00, 0x6B, 0x3F)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "EAF4EE"


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_para(doc, text, *, bold=False, size=11, color=DARK, align=None, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(sizes.get(level, 11))
    run.bold = True
    run.font.color.rgb = GREEN if level == 1 else DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        set_cell_bg(cell, "00 6B 3F".replace(" ", ""))
        set_cell_borders(cell)
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            set_cell_borders(cell)
            if r % 2 == 0:
                set_cell_bg(cell, LIGHT_BG)
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for c, w in enumerate(col_widths):
                row.cells[c].width = w
    return table


def add_header_block(doc):
    # En-tête identification documentaire
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    head = ["Référence", "Version", "Date d'application", "Rédacteur"]
    vals = ["FM-IT-XX", "01", date.today().strftime("%d/%m/%Y"), "T. HOCEDEZ — RQSE"]
    for i, h in enumerate(head):
        cell = table.cell(0, i)
        cell.text = ""
        set_cell_bg(cell, "00 6B 3F".replace(" ", ""))
        set_cell_borders(cell)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, v in enumerate(vals):
        cell = table.cell(1, i)
        cell.text = ""
        set_cell_borders(cell)
        run = cell.paragraphs[0].add_run(v)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def main():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Style normal par défaut
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Titre principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FICHE MÉTHODE")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = GREEN
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Déploiement de l'API Sellsy dans l'application SuiviPDP")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK
    run.font.name = "Calibri"

    doc.add_paragraph()

    # Bloc identification
    add_header_block(doc)
    doc.add_paragraph()

    # 1. Objet
    add_heading(doc, "1. Objet", 1)
    add_para(
        doc,
        "Décrire la procédure permettant d'activer l'intégration de l'API Sellsy v2 "
        "dans l'application SuiviPDP, afin de récupérer automatiquement les données "
        "des sociétés et des opportunités commerciales pour pré-remplir les Plans de "
        "Prévention (PDP) et les Inspections Communes Préalables (ICP).",
    )

    # 2. Domaine d'application
    add_heading(doc, "2. Domaine d'application", 1)
    add_para(
        doc,
        "Cette fiche s'applique à toute première installation, mise à jour des "
        "identifiants ou redéploiement de l'API Sellsy sur l'environnement OVH "
        "hébergeant l'application SuiviPDP.",
    )

    # 3. Acteurs
    add_heading(doc, "3. Acteurs", 1)
    add_table(
        doc,
        headers=["Rôle", "Responsabilité"],
        rows=[
            ["Administrateur Sellsy", "Création de l'application Sellsy et génération des identifiants OAuth."],
            ["Administrateur OVH / IT", "Dépôt des fichiers PHP, création du fichier de configuration sécurisé."],
            ["Responsable QSE", "Vérification fonctionnelle de la chaîne (test de connexion + recherche)."],
        ],
        col_widths=[Cm(5), Cm(11)],
    )

    # 4. Prérequis
    add_heading(doc, "4. Prérequis", 1)
    add_bullets(doc, [
        "Compte Sellsy actif avec droits d'administration sur l'espace.",
        "Accès FTP au cluster OVH (FileZilla recommandé).",
        "Application SuiviPDP déjà déployée à la racine de /www/suivipdp/.",
        "Navigateur à jour (Chrome, Edge ou Firefox récent).",
    ])

    # 5. Logigramme synthétique
    add_heading(doc, "5. Logigramme synthétique de la procédure", 1)
    add_table(
        doc,
        headers=["Étape", "Acteur", "Action", "Livrable"],
        rows=[
            ["1", "Admin Sellsy", "Créer l'application privée dans Sellsy", "client_id + client_secret"],
            ["2", "Admin OVH", "Déposer sellsy-proxy.php sur OVH", "Fichier accessible en HTTP"],
            ["3", "Admin OVH", "Créer sellsy-config.php avec les identifiants", "Configuration sécurisée"],
            ["4", "RQSE", "Tester la connexion depuis l'app", "Statut « Connexion OK »"],
            ["5", "RQSE", "Vérifier la recherche société + opportunité", "Données importées dans PDP/ICP"],
        ],
        col_widths=[Cm(1.5), Cm(3), Cm(8), Cm(4)],
    )

    # 6. Procédure détaillée
    add_heading(doc, "6. Procédure détaillée", 1)

    add_heading(doc, "6.1 Phase 1 — Côté Sellsy (≈ 5 minutes)", 2)
    add_numbered(doc, [
        "Se connecter à Sellsy avec un compte administrateur.",
        "Aller dans : Paramètres ⚙️ → Développeurs → Mes applications privées.",
        "Cliquer sur « Créer une application » et choisir le type Personal.",
        "Cocher les permissions : Sociétés (lecture), Opportunités (lecture), Contacts (lecture).",
        "Valider et noter en lieu sûr les deux identifiants : client_id et client_secret.",
    ])
    add_para(
        doc,
        "Important : le client_secret n'est affiché qu'une seule fois. Le copier "
        "immédiatement dans un gestionnaire de mots de passe (ex. KeePass).",
        size=10, color=GREY, space_after=8,
    )

    add_heading(doc, "6.2 Phase 2 — Côté OVH (≈ 10 minutes)", 2)
    add_para(doc, "Connexion FTP au cluster OVH avec FileZilla.", bold=True)
    add_table(
        doc,
        headers=["Paramètre", "Valeur"],
        rows=[
            ["Hôte", "ftp.cluster0XX.hosting.ovh.net (selon votre hébergement)"],
            ["Identifiant", "(identifiant FTP OVH)"],
            ["Mot de passe", "(mot de passe FTP OVH)"],
            ["Port", "21 (FTP) ou 22 (SFTP)"],
            ["Dossier cible", "/www/suivipdp/"],
        ],
        col_widths=[Cm(4), Cm(12)],
    )
    doc.add_paragraph()

    add_para(doc, "a) Vérifier la présence des fichiers du proxy", bold=True)
    add_para(
        doc,
        "Les fichiers sellsy-proxy.php et sellsy-config.example.php doivent être "
        "présents dans /www/suivipdp/ (ils sont déployés automatiquement par le "
        "workflow CI lors d'un push sur le dépôt GitHub).",
    )
    add_para(
        doc,
        "Si ces fichiers sont absents, les téléverser manuellement depuis le dépôt "
        "local situé dans Documents/suivipdp/.",
    )

    add_para(doc, "b) Créer le fichier sellsy-config.php", bold=True)
    add_numbered(doc, [
        "Dans FileZilla, faire un clic droit sur sellsy-config.example.php → Renommer une copie en sellsy-config.php.",
        "Éditer le fichier sellsy-config.php directement sur le serveur.",
        "Remplacer YOUR_SELLSY_CLIENT_ID et YOUR_SELLSY_CLIENT_SECRET par les valeurs notées en phase 1.",
        "Enregistrer.",
    ])
    add_para(
        doc,
        "Sécurité : ce fichier ne doit JAMAIS être committé sur GitHub. "
        "Il est volontairement exclu via le .gitignore. Il reste exclusivement "
        "sur le serveur OVH.",
        size=10, color=GREY, space_after=8,
    )

    add_para(doc, "c) Vérifier les droits", bold=True)
    add_table(
        doc,
        headers=["Élément", "Droits attendus"],
        rows=[
            ["sellsy-config.php", "644 (lecture serveur uniquement)"],
            ["Dossier /www/suivipdp/", "755 (écriture pour cache du token)"],
            ["sellsy-proxy.php", "644"],
        ],
        col_widths=[Cm(7), Cm(9)],
    )
    doc.add_paragraph()

    add_heading(doc, "6.3 Phase 3 — Configuration dans l'application (≈ 1 minute)", 2)
    add_numbered(doc, [
        "Ouvrir l'application SuiviPDP dans le navigateur.",
        "Aller dans Paramètres → section Intégration Sellsy.",
        "Champ URL du proxy : laisser vide si l'app et le proxy sont sur le même domaine, sinon saisir l'URL complète https://votre-domaine.fr/suivipdp/sellsy-proxy.php.",
        "Cliquer sur Enregistrer.",
        "Cliquer sur Tester la connexion.",
    ])
    add_para(
        doc,
        "Résultat attendu : un message vert indiquant « Connexion Sellsy OK "
        "(token obtenu) » doit s'afficher.",
        bold=True, color=GREEN,
    )

    add_heading(doc, "6.4 Phase 4 — Vérification fonctionnelle", 2)
    add_table(
        doc,
        headers=["Test", "Emplacement", "Résultat attendu"],
        rows=[
            ["Recherche société", "PDP → Entreprises → bouton « Rechercher dans Sellsy »", "Liste de sociétés Sellsy affichée"],
            ["Import société", "Cliquer sur un résultat de recherche", "EE pré-remplie : nom, SIRET, adresse, téléphone, email, responsable, activité"],
            ["Recherche opportunité", "ICP → bouton 🔍 vert à côté du champ N° d'opération", "Liste d'opportunités affichée"],
            ["Import opportunité", "Cliquer sur une opportunité", "N° opération + objet + site + représentant EU pré-remplis"],
            ["Synchronisation complète", "Paramètres → bouton « Synchroniser depuis Sellsy »", "Toutes les sociétés Sellsy importées dans le répertoire local"],
        ],
        col_widths=[Cm(4), Cm(6), Cm(6)],
    )

    # 7. Diagnostic
    add_heading(doc, "7. Diagnostic en cas de défaillance", 1)
    add_table(
        doc,
        headers=["Symptôme", "Cause probable", "Action corrective"],
        rows=[
            ["⚠️ Message « sellsy-config.php manquant ou incomplet »", "Fichier absent ou vide sur OVH", "Recopier le fichier via FileZilla et vérifier le contenu"],
            ["❌ Erreur 401 « Credentials refusés »", "Mauvais client_id ou client_secret", "Régénérer les identifiants côté Sellsy et les saisir à nouveau"],
            ["❌ « Proxy inaccessible »", "Mauvaise URL du proxy ou PHP désactivé", "Tester l'URL https://domaine/suivipdp/sellsy-proxy.php?action=healthcheck dans le navigateur"],
            ["❌ Erreur CORS dans la console", "Application servie sur un domaine différent du proxy", "Déployer le proxy sur le même domaine ou ajuster les origines autorisées dans sellsy-proxy.php (ligne 17)"],
            ["⏳ Lenteur après mise à jour des identifiants", "Cache token obsolète", "Supprimer le fichier .sellsy-token.cache sur OVH"],
        ],
        col_widths=[Cm(5.5), Cm(4.5), Cm(6)],
    )

    # 8. Sécurité
    add_heading(doc, "8. Points de vigilance sécurité", 1)
    add_bullets(doc, [
        "Le client_secret reste stocké uniquement sur le serveur OVH et n'est jamais transmis au navigateur.",
        "Le fichier sellsy-config.php est exclu du dépôt Git via le .gitignore.",
        "En production, restreindre les origines CORS du proxy à votre domaine uniquement (sellsy-proxy.php, ligne 17 : remplacer ['*'] par ['https://votre-domaine.fr']).",
        "Le cache du token (.sellsy-token.cache) est régénéré automatiquement à chaque expiration (environ une heure).",
        "Effectuer une rotation des identifiants Sellsy en cas de suspicion de compromission.",
    ])

    # 9. Documents associés
    add_heading(doc, "9. Documents et références associés", 1)
    add_bullets(doc, [
        "Code source SuiviPDP : https://github.com/Cocothib/suivipdp",
        "Documentation API Sellsy v2 : https://api.sellsy.com/doc/v2/",
        "Procédure de gestion des accès et habilitations (référence interne SMI à compléter).",
        "Politique de sécurité des systèmes d'information AGRIWATT (à compléter).",
    ])

    # 10. Validation
    add_heading(doc, "10. Validation et historique", 1)
    add_table(
        doc,
        headers=["Version", "Date", "Modification", "Rédacteur", "Approbateur"],
        rows=[
            ["01", date.today().strftime("%d/%m/%Y"), "Création initiale du document", "T. HOCEDEZ", "—"],
        ],
        col_widths=[Cm(2), Cm(3), Cm(5), Cm(3), Cm(3)],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "— Fin du document —"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    run.font.name = "Calibri"

    doc.save(DOC_PATH)
    print(f"Fichier créé : {DOC_PATH}")


if __name__ == "__main__":
    main()
