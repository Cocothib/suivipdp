"""Génère 2 supports DOCX pour la réunion avec le responsable maintenance :
1. Support de présentation A4 (recto-verso, pour la réunion)
2. Fiche réflexe technicien (1 page, plastifiable, pour le terrain)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Couleurs AGRIWATT (vert + gris)
VERT = RGBColor(0x2E, 0x7D, 0x32)
GRIS_FONCE = RGBColor(0x37, 0x47, 0x4F)
GRIS_CLAIR = RGBColor(0xEC, 0xEF, 0xF1)
ORANGE = RGBColor(0xE6, 0x5C, 0x00)
ROUGE = RGBColor(0xC6, 0x28, 0x28)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="2E7D32", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_margins(doc, top=1.5, bottom=1.5, left=1.8, right=1.8):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_titre(doc, texte, taille=20, couleur=VERT, centre=True):
    p = doc.add_paragraph()
    if centre:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.name = 'Calibri'
    run.font.size = Pt(taille)
    run.font.bold = True
    run.font.color.rgb = couleur
    return p


def add_sous_titre(doc, texte, taille=12, couleur=GRIS_FONCE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.name = 'Calibri'
    run.font.size = Pt(taille)
    run.font.italic = True
    run.font.color.rgb = couleur
    return p


def add_section(doc, numero, titre):
    p = doc.add_paragraph()
    run_num = p.add_run(f" {numero} ")
    run_num.font.bold = True
    run_num.font.size = Pt(13)
    run_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # fond vert via shading sur le paragraphe
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), '2E7D32')
    pPr.append(shd)
    run_titre = p.add_run(f"  {titre}")
    run_titre.font.bold = True
    run_titre.font.size = Pt(13)
    run_titre.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return p


def add_para(doc, texte, taille=10, gras=False, italique=False, couleur=None, alignement=None):
    p = doc.add_paragraph()
    if alignement:
        p.alignment = alignement
    run = p.add_run(texte)
    run.font.name = 'Calibri'
    run.font.size = Pt(taille)
    run.font.bold = gras
    run.font.italic = italique
    if couleur:
        run.font.color.rgb = couleur
    return p


def add_bullet(doc, texte, taille=10, gras_debut=None):
    p = doc.add_paragraph(style='List Bullet')
    if gras_debut:
        run_g = p.add_run(gras_debut)
        run_g.font.bold = True
        run_g.font.size = Pt(taille)
        run_g.font.name = 'Calibri'
        run_n = p.add_run(texte)
        run_n.font.size = Pt(taille)
        run_n.font.name = 'Calibri'
    else:
        run = p.add_run(texte)
        run.font.size = Pt(taille)
        run.font.name = 'Calibri'
    return p


# ============================================================
# DOCUMENT 1 — SUPPORT REUNION (recto-verso A4)
# ============================================================
def generer_support_reunion(chemin):
    doc = Document()
    set_margins(doc, top=1.2, bottom=1.2, left=1.5, right=1.5)

    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # --- EN-TÊTE ---
    add_titre(doc, "DÉPLOIEMENT APPLICATION SUIVI ICP / PDP", taille=18)
    add_sous_titre(doc, "Réunion responsable maintenance opérationnelle  —  AGRIWATT", taille=10)
    add_para(doc, "https://zbpbasv.cluster121.hosting.ovh.net/suivipdp/",
             taille=9, italique=True, couleur=GRIS_FONCE,
             alignement=WD_ALIGN_PARAGRAPH.CENTER)

    # --- 1. POURQUOI ---
    add_section(doc, "1", "POURQUOI CET OUTIL — LE CADRE LÉGAL")
    add_para(doc,
             "Le décret n° 92-158 impose la rédaction d'un Plan de Prévention (PDP) "
             "et la réalisation d'une Inspection Commune Préalable (ICP) avant toute "
             "intervention d'une entreprise extérieure sur un site client.",
             taille=10)
    add_para(doc, "Sans ICP signée et tracée :", taille=10, gras=True, couleur=ROUGE)
    add_bullet(doc, "responsabilité pénale du chef d'établissement engagée",
               gras_debut="• ")
    add_bullet(doc, "couverture assurance fragilisée en cas d'accident",
               gras_debut="• ")
    add_bullet(doc, "non-conformité majeure en audit ISO 9001 / 14001 / 45001",
               gras_debut="• ")
    add_bullet(doc, "perte de marchés (clients industriels = exigence contractuelle)",
               gras_debut="• ")

    # --- 2. CE QUE L'APPLI APPORTE ---
    add_section(doc, "2", "CE QUE L'APPLICATION CHANGE POUR LE TERRAIN")
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(8.5)

    donnees = [
        ("AVANT  —  flux papier / mail", "AVEC SuiviPDP  —  flux numérique"),
        ("ICP papier remplie sur site (15-20 min)", "Check-list guidée sur smartphone (8-10 min)"),
        ("Ressaisie au bureau (10-15 min)", "Aucune ressaisie  —  PDF généré automatiquement"),
        ("Archivage manuel, recherche difficile", "Archivage centralisé, recherche instantanée"),
    ]
    for i, (avant, apres) in enumerate(donnees):
        cell_a = table.cell(i, 0)
        cell_b = table.cell(i, 1)
        cell_a.text = avant
        cell_b.text = apres
        for cell in (cell_a, cell_b):
            set_cell_borders(cell)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if i == 0:
            set_cell_bg(cell_a, "37474F")
            set_cell_bg(cell_b, "2E7D32")
        elif i % 2 == 1:
            set_cell_bg(cell_a, "ECEFF1")
            set_cell_bg(cell_b, "E8F5E9")

    add_para(doc, "")

    # --- 3. PARCOURS TECHNICIEN ---
    add_section(doc, "3", "PARCOURS TECHNICIEN  —  3 ÉTAPES SUR SITE")
    add_bullet(doc,
               "Ouverture appli  →  sélection du client (base Sellsy synchronisée, pas de re-saisie)",
               gras_debut="ÉTAPE 1  —  ")
    add_bullet(doc,
               "Check-list ICP guidée  →  risques chimiques, électriques, hauteur, co-activité  →  photos jointes",
               gras_debut="ÉTAPE 2  —  ")
    add_bullet(doc,
               "Signatures électroniques EU + EE sur place  →  PDF généré  →  envoi automatique client + archivage",
               gras_debut="ÉTAPE 3  —  ")
    add_para(doc,
             "Mode hors-ligne natif : fonctionne sans réseau, synchronisation au retour.",
             taille=9, italique=True, couleur=VERT)

    # --- 4. PLAN DE DÉPLOIEMENT ---
    add_section(doc, "4", "PLAN DE DÉPLOIEMENT PROPOSÉ  —  6 SEMAINES")
    table2 = doc.add_table(rows=5, cols=3)
    table2.autofit = False
    en_tetes = ["PHASE", "DURÉE", "ACTIONS"]
    lignes = [
        ("1. Pilote",       "Semaines 1-2", "2 techniciens volontaires  •  retours quotidiens  •  ajustements appli"),
        ("2. Formation",    "Semaine 3",    "Session 1h avec l'équipe  •  remise fiche réflexe  •  désignation référent terrain"),
        ("3. Bascule",      "Semaines 4-5", "100 % des ICP sur l'appli  •  refus du papier au-delà"),
        ("4. Bilan",        "Semaine 6",    "Mesure du gain de temps  •  retours techniciens  •  ajustements finaux"),
    ]
    en_tete_cells = table2.rows[0].cells
    for i, h in enumerate(en_tetes):
        en_tete_cells[i].text = h
        set_cell_bg(en_tete_cells[i], "2E7D32")
        set_cell_borders(en_tete_cells[i])
        for para in en_tete_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for i, (phase, duree, actions) in enumerate(lignes, start=1):
        cells = table2.rows[i].cells
        cells[0].text = phase
        cells[1].text = duree
        cells[2].text = actions
        for j, cell in enumerate(cells):
            set_cell_borders(cell)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = VERT
    table2.columns[0].width = Cm(3.5)
    table2.columns[1].width = Cm(3)
    table2.columns[2].width = Cm(10.5)

    # --- 5. RÔLE RESPONSABLE MAINTENANCE ---
    add_section(doc, "5", "RÔLE DU RESPONSABLE MAINTENANCE  —  CONDITION DE RÉUSSITE")
    add_para(doc, "Sans portage hiérarchique visible, l'outil meurt en 3 mois. "
                  "Ce que je te demande concrètement :",
             taille=10, gras=True, couleur=ORANGE)
    add_bullet(doc, "annoncer toi-même le déploiement en réunion d'équipe (légitimité)",
               gras_debut="✓  ")
    add_bullet(doc, "intégrer l'usage de l'appli dans les objectifs individuels 2026",
               gras_debut="✓  ")
    add_bullet(doc, "refuser toute ICP hors-appli passée la phase pilote",
               gras_debut="✓  ")
    add_bullet(doc, "valoriser publiquement les techniciens qui jouent le jeu",
               gras_debut="✓  ")
    add_bullet(doc, "remonter en COMEX 1 indicateur mensuel : % ICP réalisées via l'appli",
               gras_debut="✓  ")

    # --- PIED ---
    add_para(doc, "")
    p_pied = doc.add_paragraph()
    p_pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_pied.add_run("T. HOCEDEZ  —  Responsable QSE AGRIWATT  —  thocedez@agriwatt.fr")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = GRIS_FONCE

    doc.save(chemin)
    print(f"OK  Support reunion :  {chemin}")


# ============================================================
# DOCUMENT 2 — FICHE RÉFLEXE TECHNICIEN (1 page plastifiable)
# ============================================================
def generer_fiche_reflexe(chemin):
    doc = Document()
    set_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- BANDEAU TITRE ---
    add_titre(doc, "FICHE RÉFLEXE  —  ICP / PDP SUR SITE", taille=20)
    add_sous_titre(doc,
                   "Application SuiviPDP  —  À garder dans le véhicule",
                   taille=11)
    add_para(doc, "")

    # --- AVANT DE PARTIR ---
    add_section(doc, "1", "AVANT DE PARTIR DU DÉPÔT")
    add_bullet(doc, "Smartphone chargé + appli SuiviPDP installée",
               taille=11, gras_debut="☐  ")
    add_bullet(doc, "Synchronisation OK (icône verte en haut à droite)",
               taille=11, gras_debut="☐  ")
    add_bullet(doc, "Client présent dans la liste  →  sinon prévenir QSE",
               taille=11, gras_debut="☐  ")

    # --- SUR SITE ---
    add_section(doc, "2", "À L'ARRIVÉE SUR SITE CLIENT")
    add_bullet(doc, "Ouvrir l'appli  →  «  Nouvelle ICP  »  →  sélectionner le client",
               taille=11, gras_debut="①  ")
    add_bullet(doc, "Identifier le représentant client (EU) qui signera avec toi",
               taille=11, gras_debut="②  ")
    add_bullet(doc,
               "Faire le tour avec lui  →  cocher chaque risque  →  photos des points sensibles",
               taille=11, gras_debut="③  ")
    add_bullet(doc, "Compléter les mesures de prévention (champs guidés)",
               taille=11, gras_debut="④  ")
    add_bullet(doc, "Signature EU + signature EE sur l'écran  →  «  Valider  »",
               taille=11, gras_debut="⑤  ")
    add_bullet(doc, "PDF envoyé automatiquement au client + archivé",
               taille=11, gras_debut="⑥  ")

    # --- POINTS DE VIGILANCE ---
    add_section(doc, "3", "POINTS DE VIGILANCE  —  NE JAMAIS OUBLIER")
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(8.5)

    vigilances = [
        ("⚠  Co-activité",          "Y a-t-il d'autres entreprises sur site ? Le noter."),
        ("⚠  Travail en hauteur",   "> 2 m  →  harnais + point d'ancrage tracés dans l'ICP."),
        ("⚠  Risque électrique",    "Consignation faite ? Par qui ? Heure ? Photo de l'attestation."),
        ("⚠  Modification en cours", "Si le chantier change, REFAIRE une ICP. Pas d'avenant verbal."),
    ]
    for i, (titre_v, detail) in enumerate(vigilances):
        cell_a = table.cell(i, 0)
        cell_b = table.cell(i, 1)
        cell_a.text = titre_v
        cell_b.text = detail
        for cell in (cell_a, cell_b):
            set_cell_borders(cell, color="E65C00", sz="8")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
        for run in cell_a.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = ORANGE
        set_cell_bg(cell_a, "FFF3E0")

    add_para(doc, "")

    # --- HORS-LIGNE ---
    add_section(doc, "4", "PAS DE RÉSEAU SUR SITE ?")
    add_para(doc,
             "Aucun problème  —  l'appli enregistre tout en local. "
             "La synchronisation se fait automatiquement dès qu'un réseau est disponible "
             "(retour véhicule, dépôt, 4G). NE PAS supprimer l'appli.",
             taille=11)

    # --- AIDE ---
    add_section(doc, "5", "BESOIN D'AIDE  —  CONTACTS")
    table_c = doc.add_table(rows=3, cols=2)
    table_c.autofit = False
    contacts = [
        ("Problème technique appli",       "Référent terrain  →  [à compléter]"),
        ("Question ICP / sécurité",        "QSE  —  T. HOCEDEZ  →  thocedez@agriwatt.fr"),
        ("Client absent / refus signer",   "Responsable maintenance  →  [à compléter]"),
    ]
    for i, (motif, contact) in enumerate(contacts):
        cell_a = table_c.cell(i, 0)
        cell_b = table_c.cell(i, 1)
        cell_a.text = motif
        cell_b.text = contact
        for cell in (cell_a, cell_b):
            set_cell_borders(cell, color="2E7D32")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
        for run in cell_a.paragraphs[0].runs:
            run.font.bold = True
        set_cell_bg(cell_a, "E8F5E9")
    table_c.columns[0].width = Cm(7)
    table_c.columns[1].width = Cm(10)

    # --- PIED ---
    add_para(doc, "")
    p_pied = doc.add_paragraph()
    p_pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_pied.add_run(
        "Pas d'ICP signée  =  pas d'intervention.  Ta sécurité, ta protection juridique."
    )
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = ROUGE

    doc.save(chemin)
    print(f"OK  Fiche reflexe :   {chemin}")


if __name__ == "__main__":
    import os
    dossier = r"C:\Users\ThibaultHOCEDEZ\Documents\suivipdp"
    generer_support_reunion(os.path.join(dossier, "Support_Reunion_ICP_PDP.docx"))
    generer_fiche_reflexe(os.path.join(dossier, "Fiche_Reflexe_Technicien_ICP.docx"))
